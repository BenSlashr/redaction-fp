import os
import uuid
import logging
from typing import Optional, Dict, Any, List
from fastapi import UploadFile

# Importation conditionnelle de aiofiles
try:
    import aiofiles
    AIOFILES_AVAILABLE = True
except ImportError:
    logger = logging.getLogger(__name__)
    logger.warning("Module aiofiles non disponible. Certaines fonctionnalités de traitement de fichiers seront limitées.")
    AIOFILES_AVAILABLE = False

# Importation conditionnelle de PyPDF2
try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    logger = logging.getLogger(__name__)
    logger.warning("Module PyPDF2 non disponible. Le traitement des fichiers PDF sera limité.")
    PYPDF2_AVAILABLE = False

# Importation conditionnelle de pdfplumber
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    logger = logging.getLogger(__name__)
    logger.warning("Module pdfplumber non disponible. Le traitement des fichiers PDF sera limité.")
    PDFPLUMBER_AVAILABLE = False

# Importation conditionnelle de docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    logger = logging.getLogger(__name__)
    logger.warning("Module docx non disponible. Le traitement des fichiers DOCX sera limité.")
    DOCX_AVAILABLE = False

# Importation conditionnelle de BeautifulSoup
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    logger = logging.getLogger(__name__)
    logger.warning("Module BeautifulSoup non disponible. Le traitement des fichiers HTML sera limité.")
    BS4_AVAILABLE = False

import io
import re
from models.client_document import ClientDocument

logger = logging.getLogger(__name__)

class FileProcessor:
    """
    Service pour traiter les fichiers téléchargés et en extraire le contenu.
    """
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.txt': 'text/plain',
        '.html': 'text/html',
        '.htm': 'text/html',
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 Mo
    
    @staticmethod
    async def _extract_content(file: UploadFile, file_extension: str) -> str:
        """
        Extrait le contenu textuel d'un fichier en fonction de son type.
        
        Args:
            file: Fichier téléchargé
            file_extension: Extension du fichier
            
        Returns:
            str: Contenu textuel extrait du fichier
        """
        content = ""
        
        # Lire le contenu du fichier
        file_content = await file.read()
        
        # Extraire le texte en fonction du type de fichier
        if file_extension.lower() in ['.pdf']:
            # Vérifier si les modules PDF sont disponibles
            if not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
                logger.error("Aucun module de traitement PDF n'est disponible. Impossible d'extraire le texte du PDF.")
                content = "[Contenu PDF non extractible - modules manquants]\n"
            else:
                # Essayer d'abord avec pdfplumber si disponible
                if PDFPLUMBER_AVAILABLE:
                    try:
                        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                            for page in pdf.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    content += page_text + "\n\n"
                    except Exception as e:
                        logger.error(f"Erreur lors de l'extraction du texte du PDF avec pdfplumber: {str(e)}")
                
                # Si pdfplumber n'a pas extrait de texte ou n'est pas disponible, essayer avec PyPDF2
                if (not content.strip() or not PDFPLUMBER_AVAILABLE) and PYPDF2_AVAILABLE:
                    try:
                        pdf_reader = PdfReader(io.BytesIO(file_content))
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                content += page_text + "\n\n"
                    except Exception as e:
                        logger.error(f"Erreur lors de l'extraction du texte du PDF avec PyPDF2: {str(e)}")
                    
        elif file_extension.lower() in ['.docx', '.doc']:
            # Vérifier si le module docx est disponible
            if not DOCX_AVAILABLE:
                logger.error("Le module docx n'est pas disponible. Impossible d'extraire le texte du document Word.")
                content = "[Contenu Word non extractible - module manquant]\n"
            else:
                # Traitement des documents Word
                try:
                    doc = Document(io.BytesIO(file_content))
                    for para in doc.paragraphs:
                        if para.text:
                            content += para.text + "\n"
                except Exception as e:
                    logger.error(f"Erreur lors de l'extraction du texte du document Word: {str(e)}")
                    content = "[Erreur lors de l'extraction du contenu Word]\n"
                    
        elif file_extension.lower() in ['.txt']:
            # Traitement des fichiers texte
            content = file_content.decode('utf-8', errors='replace')
            
        elif file_extension.lower() in ['.html', '.htm']:
            # Vérifier si le module BeautifulSoup est disponible
            if not BS4_AVAILABLE:
                logger.error("Le module BeautifulSoup n'est pas disponible. Impossible d'extraire le texte du fichier HTML.")
                content = "[Contenu HTML non extractible - module manquant]\n"
            else:
                # Traitement des fichiers HTML
                try:
                    soup = BeautifulSoup(file_content, 'html.parser')
                    
                    # Supprimer les scripts, styles et balises de commentaires
                    for script in soup(["script", "style"]):
                        script.extract()
                        
                    # Extraire le texte
                    content = soup.get_text(separator="\n")
                    
                    # Nettoyer les espaces et lignes vides multiples
                    content = re.sub(r'\n\s*\n', '\n\n', content)
                except Exception as e:
                    logger.error(f"Erreur lors de l'extraction du texte du fichier HTML: {str(e)}")
                    content = "[Erreur lors de l'extraction du contenu HTML]\n"
            
        else:
            raise ValueError(f"Type de fichier non pris en charge: {file_extension}")
            
        # Nettoyer le contenu
        content = content.strip()
        
        # Réinitialiser la position du fichier pour une utilisation ultérieure
        await file.seek(0)
        
        return content
    
    @staticmethod
    def _get_file_extension(filename: str) -> str:
        """
        Récupère l'extension d'un fichier.
        
        Args:
            filename: Nom du fichier
            
        Returns:
            str: Extension du fichier (avec le point)
        """
        _, ext = os.path.splitext(filename)
        return ext.lower()
    
    @staticmethod
    def _validate_file(file: UploadFile) -> str:
        """
        Valide le fichier téléchargé (taille, type, etc.).
        
        Args:
            file: Fichier téléchargé
            
        Returns:
            str: Extension du fichier validé
            
        Raises:
            ValueError: Si le fichier n'est pas valide
        """
        # Vérifier que le fichier existe
        if not file or not file.filename:
            raise ValueError("Aucun fichier fourni")
            
        # Récupérer l'extension du fichier
        file_extension = FileProcessor._get_file_extension(file.filename)
        
        # Vérifier que le type de fichier est pris en charge
        if file_extension not in FileProcessor.SUPPORTED_EXTENSIONS:
            supported_exts = ", ".join(FileProcessor.SUPPORTED_EXTENSIONS.keys())
            raise ValueError(f"Type de fichier non pris en charge. Types acceptés: {supported_exts}")
            
        # Vérifier que le content_type correspond à l'extension
        expected_content_type = FileProcessor.SUPPORTED_EXTENSIONS[file_extension]
        if file.content_type and not (
            file.content_type == expected_content_type or 
            file.content_type == 'application/octet-stream'  # Fallback pour les types non détectés
        ):
            raise ValueError(f"Le type de contenu du fichier ({file.content_type}) ne correspond pas à l'extension ({file_extension})")
            
        return file_extension
    
    async def process_uploaded_file(
        self, 
        file: UploadFile, 
        client_id: str, 
        title: Optional[str] = None, 
        source_type: str = "uploaded_file"
    ) -> ClientDocument:
        """
        Traite un fichier téléchargé et en extrait le contenu.
        
        Args:
            file: Fichier téléchargé
            client_id: ID du client
            title: Titre du document (facultatif, utilisera le nom du fichier si non fourni)
            source_type: Type de source du document
            
        Returns:
            ClientDocument: Document client créé à partir du fichier
            
        Raises:
            ValueError: Si le fichier n'est pas valide
        """
        try:
            # Valider le fichier
            file_extension = self._validate_file(file)
            
            # Utiliser le nom du fichier comme titre si non fourni
            if not title:
                title = os.path.splitext(file.filename)[0]
                
            # Extraire le contenu du fichier
            content = await self._extract_content(file, file_extension)
            
            # Créer un ID unique pour le document
            document_id = str(uuid.uuid4())
            
            # Créer le document client
            document = ClientDocument(
                document_id=document_id,
                client_id=client_id,
                title=title,
                content=content,
                source_type=source_type,
                metadata={
                    "filename": file.filename,
                    "file_extension": file_extension,
                    "content_type": file.content_type,
                }
            )
            
            logger.info(f"Fichier '{file.filename}' traité avec succès pour le client {client_id}")
            
            return document
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement du fichier '{file.filename if file else 'inconnu'}': {str(e)}")
            raise
